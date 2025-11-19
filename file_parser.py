import os
import re
import logging
from docx import Document
import PyPDF2

logger = logging.getLogger(__name__)

class FileParser:
    @staticmethod
    def parse_file(file_path: str):
        """Fayldan savollarni o'qish - asosiy metod"""
        if not os.path.exists(file_path):
            logger.error(f"Fayl topilmadi: {file_path}")
            return []
        
        try:
            if file_path.endswith('.docx'):
                questions = FileParser.parse_docx(file_path)
                if not questions:
                    questions = FileParser.parse_docx_advanced(file_path)
                return questions
            elif file_path.endswith('.pdf'):
                questions = FileParser.parse_pdf(file_path)
                return questions
            else:
                logger.error(f"Noto'g'ri fayl formati: {file_path}")
                return []
        except Exception as e:
            logger.error(f"Faylni o'qishda xatolik: {e}")
            return []
    
    @staticmethod
    def parse_docx(file_path: str):
        """Word faylidan savollarni o'qish - asosiy metod"""
        try:
            doc = Document(file_path)
            questions = []
            current_question = None
            current_options = []
            question_number = 0
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue
                
                # Savolni aniqlash (raqam bilan boshlanadigan)
                if FileParser.is_question_start(text):
                    # Oldingi savolni saqlash
                    if current_question and len(current_options) >= 2:
                        current_question['options'] = current_options
                        questions.append(current_question)
                        question_number += 1
                    
                    # Yangi savolni boshlash
                    current_question = {
                        'question': text,
                        'options': [],
                        'correct_answer': 0  # Birinchi variant to'g'ri javob
                    }
                    current_options = []
                
                # Variantlarni aniqlash
                elif FileParser.is_option(text):
                    if current_question:
                        clean_option = FileParser.clean_option_text(text)
                        if clean_option:
                            current_options.append(clean_option)
                
                # Agar matn uzun bo'lsa va variant bo'lmasa, savol qismiga qo'shish
                elif current_question and len(current_options) == 0 and len(text) > 5:
                    current_question['question'] += " " + text
            
            # Oxirgi savolni qo'shish
            if current_question and len(current_options) >= 2:
                current_question['options'] = current_options
                questions.append(current_question)
            
            logger.info(f"Word fayldan {len(questions)} ta savol topildi")
            return questions
        
        except Exception as e:
            logger.error(f"Word faylni o'qishda xatolik: {e}")
            return []
    
    @staticmethod
    def parse_docx_advanced(file_path: str):
        """Kengaytirilgan Word faylini o'qish - muqobil metod"""
        try:
            doc = Document(file_path)
            full_text = ""
            
            # Barcha paragraflarni bitta matnga birlashtirish
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    full_text += text + "\n"
            
            return FileParser.parse_text(full_text)
        
        except Exception as e:
            logger.error(f"Kengaytirilgan Word o'qishda xatolik: {e}")
            return []
    
    @staticmethod
    def parse_pdf(file_path: str):
        """PDF faylidan savollarni o'qish"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                full_text = ""
                
                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    if text:
                        full_text += f"--- Page {page_num + 1} ---\n{text}\n"
                
                logger.info(f"PDF dan {len(full_text)} belgi o'qildi")
                return FileParser.parse_text(full_text)
        
        except Exception as e:
            logger.error(f"PDF o'qishda xatolik: {e}")
            return []
    
    @staticmethod
    def parse_text(text: str):
        """Matndan savollarni ajratib olish"""
        questions = []
        lines = text.split('\n')
        current_question = None
        current_options = []
        question_number = 0
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                i += 1
                continue
            
            # Savolni aniqlash
            if FileParser.is_question_start(line):
                # Oldingi savolni saqlash
                if current_question and len(current_options) >= 2:
                    current_question['options'] = current_options
                    questions.append(current_question)
                    question_number += 1
                
                # Yangi savolni boshlash
                current_question = {
                    'question': line,
                    'options': [],
                    'correct_answer': 0
                }
                current_options = []
            
            # Variantlarni aniqlash
            elif FileParser.is_option(line):
                if current_question:
                    clean_option = FileParser.clean_option_text(line)
                    if clean_option:
                        current_options.append(clean_option)
            
            # Keyingi qator savolning davomi bo'lishi mumkin
            elif current_question and not current_options and len(line) > 5:
                # Yangi qator bo'sh joy bilan qo'shiladi
                current_question['question'] += " " + line
            
            # Variantlarning davomi (ko'p qatorli variantlar)
            elif current_question and current_options and len(line) > 3 and not FileParser.is_option(line):
                # Oxirgi variantga qo'shish
                if current_options:
                    current_options[-1] += " " + line
            
            i += 1
        
        # Oxirgi savolni qo'shish
        if current_question and len(current_options) >= 2:
            current_question['options'] = current_options
            questions.append(current_question)
        
        logger.info(f"Matndan {len(questions)} ta savol topildi")
        return questions
    
    @staticmethod
    def is_question_start(text: str) -> bool:
        """Matn savol boshlanishi ekanligini tekshirish"""
        if not text or len(text) < 3:
            return False
        
        # Raqam bilan boshlanadigan va nuqta/qavs borligini tekshirish
        patterns = [
            r'^\d+[\.\)]',          # 1. yoki 1)
            r'^\d+\.\s',            # 1. (bo'sh joy bilan)
            r'^\d+\)\s',            # 1) (bo'sh joy bilan)
            r'^«',                  # « bilan boshlanadigan
            r'^\d+\s*\.',           # 1 . (ortiqcha bo'sh joy bilan)
        ]
        
        for pattern in patterns:
            if re.match(pattern, text):
                return True
        
        # Agar raqam bilan boshlansa va uzunligi 5 dan kichik bo'lsa
        if text[0].isdigit() and len(text) < 100:
            return True
        
        return False
    
    @staticmethod
    def is_option(text: str) -> bool:
        """Matn variant ekanligini tekshirish"""
        if not text or len(text) < 2:
            return False
        
        # Variant formatlarini tekshirish
        option_patterns = [
            r'^[A-Da-d][\.\)]',     # A), B), C), D) yoki a), b), c), d)
            r'^[A-Da-d]\.',         # A., B., C., D. yoki a., b., c., d.
            r'^[A-Da-d]\s',         # A , B , C , D  (bo'sh joy bilan)
        ]
        
        for pattern in option_patterns:
            if re.match(pattern, text):
                return True
        
        return False
    
    @staticmethod
    def clean_option_text(text: str) -> str:
        """Variant matnini tozalash"""
        if not text:
            return ""
        
        # Variant prefiksini olib tashlash
        patterns_to_remove = [
            r'^[A-Da-d][\.\)]\s*',  # A) , B. , c) , d.
            r'^[A-Da-d]\s+',        # A , B , C , D
        ]
        
        for pattern in patterns_to_remove:
            text = re.sub(pattern, '', text).strip()
        
        # Qo'shimcha tozalash
        text = text.strip()
        
        # Agar matn juda qisqa bo'lsa, e'tiborga olmaslik
        if len(text) < 2:
            return ""
        
        return text
    
    @staticmethod
    def validate_questions(questions: list) -> list:
        """Savollarni tekshirish va tozalash"""
        validated = []
        
        for i, q in enumerate(questions):
            try:
                # Asosiy maydonlarni tekshirish
                if not q.get('question') or not q.get('options'):
                    continue
                
                # Savol matnini tozalash
                question_text = q['question'].strip()
                if not question_text:
                    continue
                
                # Variantlarni tozalash
                options = []
                for opt in q['options']:
                    if isinstance(opt, str):
                        clean_opt = opt.strip()
                        if clean_opt and len(clean_opt) >= 1:
                            options.append(clean_opt)
                
                # Kamida 2 ta variant bo'lishi kerak
                if len(options) < 2:
                    continue
                
                # To'g'ri javob indeksini tekshirish
                correct_answer = q.get('correct_answer', 0)
                if correct_answer >= len(options):
                    correct_answer = 0
                
                validated.append({
                    'question': question_text,
                    'options': options,
                    'correct_answer': correct_answer
                })
                
            except Exception as e:
                logger.warning(f"Savol {i} ni tekshirishda xatolik: {e}")
                continue
        
        logger.info(f"Tekshiruvdan {len(validated)}/{len(questions)} ta savol o'tdi")
        return validated
    
    @staticmethod
    def debug_file_content(file_path: str):
        """Fayl tarkibini debug qilish"""
        try:
            if file_path.endswith('.docx'):
                doc = Document(file_path)
                content = []
                for i, paragraph in enumerate(doc.paragraphs):
                    text = paragraph.text.strip()
                    if text:
                        content.append(f"{i}: {text}")
                return "\n".join(content[:50])  # Faqat birinchi 50 qator
            else:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read(2000)  # Faqat birinchi 2000 belgi
        except Exception as e:
            return f"Debug xatosi: {e}"
    
    @staticmethod
    def analyze_file_structure(file_path: str):
        """Fayl strukturasi tahlili"""
        try:
            if file_path.endswith('.docx'):
                doc = Document(file_path)
                stats = {
                    'total_paragraphs': 0,
                    'non_empty_paragraphs': 0,
                    'question_like': 0,
                    'option_like': 0,
                    'sample_lines': []
                }
                
                for i, paragraph in enumerate(doc.paragraphs[:100]):  # Birinchi 100 paragraf
                    text = paragraph.text.strip()
                    stats['total_paragraphs'] += 1
                    
                    if text:
                        stats['non_empty_paragraphs'] += 1
                        
                        if FileParser.is_question_start(text):
                            stats['question_like'] += 1
                        
                        if FileParser.is_option(text):
                            stats['option_like'] += 1
                        
                        if len(stats['sample_lines']) < 10:
                            stats['sample_lines'].append(f"{i}: {text}")
                
                return stats
            else:
                return {"error": "Faqat .docx fayllar tahlil qilinadi"}
        
        except Exception as e:
            return {"error": str(e)}
    
    @staticmethod
    def extract_questions_with_regex(text: str):
        """Regex yordamida savollarni ajratib olish"""
        questions = []
        
        # Savol va variantlarni topish uchun pattern
        pattern = r'(\d+[\.\)]\s*[^\n]+(?:\s+[A-Da-d][\.\)][^\n]+){2,})'
        matches = re.findall(pattern, text, re.MULTILINE | re.DOTALL)
        
        for match in matches:
            try:
                lines = match.strip().split('\n')
                if len(lines) < 3:  # Kamida savol + 2 variant
                    continue
                
                question_text = lines[0].strip()
                options = []
                
                for line in lines[1:]:
                    line = line.strip()
                    if FileParser.is_option(line):
                        clean_opt = FileParser.clean_option_text(line)
                        if clean_opt:
                            options.append(clean_opt)
                
                if len(options) >= 2:
                    questions.append({
                        'question': question_text,
                        'options': options,
                        'correct_answer': 0
                    })
            
            except Exception as e:
                logger.warning(f"Regex bilan savol ajratishda xatolik: {e}")
                continue
        
        return questions